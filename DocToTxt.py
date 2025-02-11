from docx import Document

# Load the .docx file
doc = Document('HW1 Introduction 2.docx')
#replace the file name with the file you want to convert

# Extract and print the text
option = int(input("1. Print in Console \n 2. Export to Txt file "))


#conditional statement

if option == 2:
    exportName = input("Enter the name of the file you want to export: ")
    with open(exportName, 'w') as file:
        for paragraph in doc.paragraphs:
            file.write(paragraph.text + '\n')
    print("File has been exported successfully")
else:
    for paragraph in doc.paragraphs:
        print(paragraph.text)
        
#end of code
